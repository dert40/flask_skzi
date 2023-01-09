from urllib.parse import ParseResult

from flask import Flask, render_template, url_for, request, redirect, flash, send_file, abort, make_response
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import aliased, join
from sqlalchemy import select, update, delete, and_, or_
from datetime import datetime
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash, safe_join
from flask_migrate import Migrate
# from flask import flash
from sqlalchemy.dialects.postgresql import UUID
import uuid
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Alignment

from flask_admin import Admin, AdminIndexView
# from flask.contrib.sqla import ModelView

from flask_admin.contrib.sqla import ModelView

from flask_security import UserMixin, RoleMixin, SQLAlchemyUserDatastore, Security


# Для blueprint Admin
# from admin.admin import admin

# from Python_doc import files

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:1@localhost/blog'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = 'fdgdftyujnkrvh5h7k9o'

app.config['SECURITY_PASSWORD_SALT'] = 'salt'
app.config['SECURITY_PASSWORD_HASH'] = 'sha512_crypt'

UPLOAD_FOLDER = 'static/uploads/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Для blueprint Admin
# app.register_blueprint(admin, url_prefix='/admin')


db = SQLAlchemy(app)
migrate = Migrate(app, db)



login_manager = LoginManager(app)


# При создании таблиц нужно обратить внимание на присутствии в определении класс сл.: (db.Model):
class T_counter(db.Model):
    __tablename__ = 't_counter'
    id = db.Column(db.Integer, primary_key=True)
    counter = db.Column(db.Integer)
    pref = db.Column(db.String(10))

    div_id = db.Column(db.Integer, db.ForeignKey('t_division.id'))

    def __repr__(self):
        return f"<T_counter {self.id}>"


class T_skzi(db.Model):
    __tablename__ = 't_skzi'
    id = db.Column(db.Integer, primary_key=True) # - 1
    n_account = db.Column(db.String(20)) # - учетный номер - 2

    n_serial = db.Column(db.String(255)) # - Серийные номера СКЗИ - 4
    n_instances = db.Column(db.String(255))# Номера экземпляров (копий) - 5
    sopr_pis = db.Column(db.String(255))# Дата и номер сопроводительного письма - 7
    date_rec = db.Column(db.DateTime, default=datetime.now) # Дата получения - 9
    date_con = db.Column(db.DateTime, default=datetime.now)# Дата подключения (установки) -11
    n_hardware = db.Column(db.String(255))# Инвентарный (серийный) номер аппаратных средств -12
    data_destr = db.Column(db.DateTime, default=datetime.now)# Дата изъятия (уничтожения)- 13
    n_destr = db.Column(db.String(255)) # - Номер акта или расписка об уничтожении - 15
    note = db.Column(db.String(300)) # - Примечание - 16

    name_skzi = db.Column(db.Integer, db.ForeignKey('name_skzi.id')) # key for - Наименование криптосредства - 3
    div_sours = db.Column(db.Integer, db.ForeignKey('t_division.id')) # key for От кого получены - 6
    fio_owner = db.Column(db.Integer, db.ForeignKey('users.id')) # key for Ф.И.О. владельца криптосредства - 8
    fio_instal = db.Column(db.Integer, db.ForeignKey('users.id')) # key for Ф.И.О. производившего подключение (установку) -10
    fio_destr = db.Column(db.Integer, db.ForeignKey('users.id')) # key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14

    def __repr__(self):
        return f"<T_SKZI {self.id}>"


class Journal_skzi(db.Model):
    __tablename__ = 'journal_skzi'
    id = db.Column(db.Integer, primary_key=True)
    date_received = db.Column(db.DateTime, default=datetime.now)
    comment = db.Column(db.String(255))
    num_doc = db.Column(db.Integer, default=0)
    uuid = db.Column(UUID(as_uuid=True))
    data_begin = db.Column(db.DateTime, default=datetime.now)
    test = db.Column(db.String(255))

    DIV_SOUR_ID = db.Column(db.Integer, db.ForeignKey('t_division.id'))
    DIVISION_ID = db.Column(db.Integer, db.ForeignKey('t_division.id'))
    user_sour = db.Column(db.Integer, db.ForeignKey('profiles.id'))
    user_rec = db.Column(db.Integer, db.ForeignKey('profiles.id'))
    name_skzi = db.Column(db.Integer, db.ForeignKey('name_skzi.id'))
    LICENSE_ID = db.Column(db.Integer, db.ForeignKey('T_license.id'))

    # pr = db.relationship('T_division', lazy="select")

    # pr = db.relationship('T_division', backref='Journal_skzi', uselist=False)

    def __repr__(self):
        return f"<Journalskzi {self.id}>"


class Name_SKZI(db.Model):
    __tablename__ = 'name_skzi'
    id = db.Column(db.Integer, primary_key=True)
    n_skzi = db.Column(db.String(255))

    def __repr__(self):
        return f"<TS_SYSTEM {self.id}>"


class T_EVENTS(db.Model):
    __tablename__ = 't_events'
    id = db.Column(db.Integer, primary_key=True)
    Type_Events = db.Column(db.String(255))

    def __repr__(self):
        return f"<T_EVENTS {self.id}>"


class T_log(db.Model):
    __tablename__ = 't_log'
    id = db.Column(db.Integer, primary_key=True)
    COMMENT = db.Column(db.String(255))
    USER_ID = db.Column(db.Integer)
    LICENSE_ID = db.Column(db.Integer)
    data_event = db.Column(db.DateTime, default=datetime.now)

    # USER_ID = db.Column(db.Integer, db.ForeignKey('users.id'))
    # LICENSE_ID = db.Column(db.Integer, db.ForeignKey('T_license.id'))
    EVENT_TYPE = db.Column(db.Integer, db.ForeignKey('t_events.id'))


class TS_SYSTEM(db.Model):
    __tablename__ = 'ts_system'
    id = db.Column(db.Integer, primary_key=True)
    SYSTEM_NAME = db.Column(db.String(255))

    DIVISION_ID = db.Column(db.Integer, db.ForeignKey('t_division.id'))

    def __repr__(self):
        return f"<TS_SYSTEM {self.id}>"


class TC_LICENSE_SYSTEM(db.Model):
    __tablename__ = 'tc_license_system'
    id = db.Column(db.Integer, primary_key=True)

    SYSTEM_ID = db.Column(db.Integer, db.ForeignKey('ts_system.id'))
    LICENSE_ID = db.Column(db.Integer, db.ForeignKey('T_license.id'))

    def __repr__(self):
        return f"<TC_LICENSE_SYSTEM {self.id}>"


class TS_CENTER(db.Model):
    __tablename__ = 'ts_center'
    id = db.Column(db.Integer, primary_key=True)
    CERTIFICATION_CENTER = db.Column(db.String(255))

    def __repr__(self):
        return f"<TS_CENTER {self.id}>"
# class TC_LICENSE_SYSTEM(db.Model):
#     __tablename__ = 'tc_license_system'
#     id = db.Column(db.Integer, primary_key=True)
#     CERTIFICATION_CENTER = db.Column(db.String(255))

    # def __repr__(self):
    #     return f"<TC_LICENSE_SYSTEM {self.id}>"


class T_Key(db.Model):
    __tablename__ = 't_key'
    id = db.Column(db.Integer, primary_key=True)
    # ID_LICENSE = db.Column(db.Integer)
    # ID_CENTER = db.Column(db.Integer)
    DATE_STOP = db.Column(db.DateTime, default=datetime.now)
    GOST = db.Column(db.String(255))
    # KEY_OWNER_ID = db.Column(db.Integer)
    COMMENT = db.Column(db.String(255))

    ID_CENTER = db.Column(db.Integer, db.ForeignKey('ts_center.id'))
    KEY_OWNER_ID = db.Column(db.Integer, db.ForeignKey('t_division.id'))
    ID_LICENSE = db.Column(db.Integer, db.ForeignKey('T_license.id'))


    def __repr__(self):
        return f"<T_Key {self.id}>"


class Post(db.Model):
    __tablename__ = 'posts'
    id = db.Column(db.Integer, primary_key=True)
    POST_NAME = db.Column(db.String(500), nullable=False)

    def __repr__(self):
        return f"<Post {self.id}>"


class T_division(db.Model):
    __tablename__ = 't_division'
    id = db.Column(db.Integer, primary_key=True)
    division_name = db.Column(db.String(255), nullable=False)
    LICENSE_OWNER = db.Column(db.Integer)

    # pr = db.relationship('Journal_skzi', backref='T_division', uselist=False)

    def __repr__(self):
        return f"<T_division {self.id}>"


class T_license(db.Model):
    # __tablename__ = 't_license'
    id = db.Column(db.Integer, primary_key=True)
    fio = db.Column(db.String(255), nullable=False)
    # post_id = db.Column(db.Integer)
    # division_id = db.Column(db.Integer)
    computer_name = db.Column(db.String(50), nullable=False)
    license_code = db.Column(db.String(29), nullable=False)
    # date_start = db.Column(db.DateTime, default=datetime.utcnow)
    # date_stop = db.Column(db.DateTime, default=datetime.utcnow)
    date_start = db.Column(db.DateTime, default=datetime.now)
    date_stop = db.Column(db.DateTime, default=datetime.now)

    last_editor = db.Column(db.Integer)
    contract = db.Column(db.String(50), nullable=False)
    comment = db.Column(db.String(255), nullable=False)
    code = db.Column(db.String(29), nullable=False)
    Сryptosource = db.Column(db.Integer)
    author_creator_id = db.Column(db.Integer)
    a_id = db.Column(db.Integer)

    div_id = db.Column(db.Integer, db.ForeignKey('t_division.id'))
    post_id = db.Column(db.Integer, db.ForeignKey('posts.id'))
    license_owner_id = db.Column(db.Integer, db.ForeignKey('t_division.id'))

    def __repr__(self):
        return '<T_license %r>' % self.id


class Article(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    intro = db.Column(db.String(300), nullable=False)
    text = db.Column(db.Text, nullable=False)
    date = db.Column(db.DateTime, default=datetime.utcnow)

    def __repr__(self):  # При ыборе объекта будет выдаваться его id
        return '<Article %r>' % self.id


roles_users = db.Table('roles_users',
                       db.Column('user_id', db.Integer(),  db.ForeignKey('users.id')),
                       db.Column('role_id', db.Integer(),  db.ForeignKey('role.id'))
                       )


class Users(db.Model, UserMixin):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(50), unique=True)
    psw = db.Column(db.String(500), nullable=True)
    password = db.Column(db.String(500), nullable=True)
    date = db.Column(db.DateTime, default=datetime.now)
    PRUZ = db.Column(db.Boolean(), default=False)
    active = db.Column(db.Boolean(), default=False)
    roles = db.relationship('Role', secondary=roles_users, backref=db.backref('users', lazy='dynamic'))

    def __repr__(self):
        return f"<users {self.id}>"


class Role(db.Model, RoleMixin):
    __tablename__ = 'role'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True)
    description = db.Column(db.String(255))


class Profiles(db.Model):
    __tablename__ = 'profiles'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=True)
    old = db.Column(db.Integer)
    city = db.Column(db.String(100))
    PRUZ = db.Column(db.Boolean(), default=False)

    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    DIVISION_OWNER = db.Column(db.Integer, db.ForeignKey('t_division.id'))
    post_id = db.Column(db.Integer, db.ForeignKey('posts.id'))

    def __repr__(self):
        return f"<profiles {self.id}>"


class AdminMixin:
    def is_accessible(self):
        return current_user.has_role('admin')

    def inaccessible_callback(self, name, **kwargs):
        return redirect( url_for('security.login', next=request.url))


# класс проверяет доступ к Admin по условию
class AdminView(AdminMixin, ModelView):
    pass
    # def is_accessible(self):
    #     return current_user.has_role('admin') # пользователь имеет роль админ
    #
    # def inaccessible_callback(self, name, **kwargs): # иначе срабатывает если вьюха не достпна текущ. полбз
    #     return redirect( url_for('security.login', next=request.url))


class HomeAdminView(AdminMixin, AdminIndexView):
    pass
    # def is_accessible(self):
    #     return current_user.has_role('admin')
    #
    # def inaccessible_callback(self, name, **kwargs):
    #     return redirect(url_for('security.login', next=request.url))

# === ADMIN ===
admin = Admin(app, 'FlasApp', url='/', index_view=HomeAdminView(name='Home'))

admin.add_view(AdminView(Users, db.session, name='Пользователи'))
admin.add_view(ModelView(Profiles, db.session, name='Профиль'))
admin.add_view(ModelView(Name_SKZI, db.session, name='Наименование криптосредства'))
admin.add_view(ModelView(T_division, db.session, name='Подразделение'))
admin.add_view(ModelView(Post, db.session, name='Должность'))
admin.add_view(ModelView(TS_CENTER, db.session, name='Удостоверяющий центр'))  #TC_LICENSE_SYSTEM
admin.add_view(ModelView(TS_SYSTEM, db.session, name='Инф. система'))

#---
user_datastore = SQLAlchemyUserDatastore(db, Users, Role)
security = Security(app, user_datastore)


#---глобальная переменная для выделенной строки


selStr = 0

#--------
# def func():
#     res = make_response("Setting a cookie")
#     res.set_cookie('foo2', 'bar2', max_age=60*60*24*365*2)
#     print('Go!!!')
#     return res


def validate(date_text):
    try:
        if date_text != datetime.strptime(date_text, "%Y-%m-%d").strftime('%Y-%m-%d'):
            raise ValueError
        return True
    except ValueError:
        return False


# Одна функия отслеживает несколько страниц

@login_manager.user_loader
def loader_user(user_id):
    return Users.query.get(user_id)


# @app.route('/')
# @app.route('/home')
# @login_required
# def index():
#     return render_template("license.html")  # какой шаблон я хочу вывести


# @app.route("/register", methods=("POST", "GET"))
# def register():
#     if request.method == "POST":
#         try:
#            hash = generate_password_hash(request.form['psw'])
#            u = Users(email=request.form['email'], psw=hash)
#            db.session.add(u)
#            db.session.flush()
#
#            p = Profiles(name=request.form['name'], old=request.form['old'], city=request.form['city'], user_id=u.id)
#            db.session.add(p)
#            db.session.commit()
#
#         except:
#             db.session.rollback()
#             print("Ошибка добавления в БД")
#
#     return render_template("register.html", title="Регистрация")
# Load Browser Favorite Icon


# @app.route('/admin/', methods=("POST", "GET"))
# @login_required
# def admin():
#         if request.method == "POST":
#             return "---"
#         else:
#             return "!!!"
#
#         return render_template("admin/index.html")


@app.route('/cookie/')
def cookie():
    res = make_response("Setting a cookie")
    res.set_cookie('foo', 'bar', max_age=60*60*24*365*2)
    return res


@app.route('/favicon.ico')
def favicon():
    return url_for('static', filename='img/favicon.ico')


@app.route('/about')
@login_required
def about():
    return render_template("about.html")


@app.route('/posts')
@login_required
def posts():
    articles = Article.query.order_by(Article.date.desc()).all()
    return render_template("posts.html", articles=articles)


@app.route('/license')
@login_required
def license():
    # spisok_lic = T_license.query.order_by(T_license.fio.desc()).all()
    q = request.args.get('q')
    search = "%{}%".format(q)

    a = aliased(T_division)
    if q:
        # spisok_lic = db.session.query(T_license, T_division, a, Post).filter(T_division.id == T_license.div_id).filter(
        #     a.id == T_license.license_owner_id).filter(Post.id == T_license.post_id).filter(T_license.fio.like(search)).all()
        spisok_lic = db.session.query(T_license, T_division, a, Post).filter(T_division.id == T_license.div_id).filter(
            a.id == T_license.license_owner_id).filter(Post.id == T_license.post_id).filter(T_license.fio.like(search) | T_division.division_name.like(search) | a.division_name.like(search)).all()
    else:
        spisok_lic = db.session.query(T_license, T_division, a, Post).filter(T_division.id == T_license.div_id).filter(
             a.id == T_license.license_owner_id).filter(Post.id == T_license.post_id).order_by(T_license.fio.desc()).all()

    return render_template("license.html", spisok_lic=spisok_lic)


@app.route('/')
@app.route('/home')
@app.route('/list')
@app.route('/list/')
@login_required
def list():
    # sp = T_skzi.query.order_by(T_skzi.id.asc()).all()
    global selStr
    str_id = selStr
    selStr = 0
    err = 'Неверно указан период дат!'
    f= request.args.get('f')

    q = request.args.get('q')

    q1 = request.args.get('q1')
    q2 = request.args.get('q2')
    q3 = request.args.get('q3')
    q4 = request.args.get('q4')
    q5 = request.args.get('q5')
    q6 = request.args.get('q6')
    q7 = request.args.get('q7')
    # q8 = request.args.get('q8')
    q8_1 =  request.args.get('q8_1')
    q8_2 =  request.args.get('q8_2')
    q9 = request.args.get('q9')

    q10_1 = request.args.get('q10_1')
    q10_2 = request.args.get('q10_2')

    q11 = request.args.get('q11')
    q12 = request.args.get('q12')

    q12_1 = request.args.get('q12_1')
    q12_2 = request.args.get('q12_2')

    q13 = request.args.get('q13')
    q14 = request.args.get('q14')
    q15 = request.args.get('q15')
    q16 = request.args.get('q16')

    search = "%{}%".format(q)
    search_1 = "%{}%".format(q1)
    search_2 = "%{}%".format(q2)
    search_3 = "%{}%".format(q3)
    search_4 = "%{}%".format(q4)
    search_5 = "%{}%".format(q5)
    search_6 = "%{}%".format(q6)
    search_7 = "%{}%".format(q7)
    # search_8 = "%{}%".format(q8)
    search_9 = "%{}%".format(q9)
    # search_10 = "%{}%".format(q10)
    search_11 = "%{}%".format(q11)
    search_12 = "%{}%".format(q12)
    search_13 = "%{}%".format(q13)
    search_14 = "%{}%".format(q14)
    search_15 = "%{}%".format(q15)
    search_16 = "%{}%".format(q16)


    users_o = aliased(Users)
    users_i = aliased(Users)
    users_d = aliased(Users)
    prof_o = aliased(Profiles)
    prof_i = aliased(Profiles)
    prof_d = aliased(Profiles)

    log = ""
    if request.cookies.get('q2'):
        log = request.cookies.get('q2')

    if q or f:


        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(T_skzi.n_account.like(search) |T_skzi.n_serial.like(search) | T_skzi.n_instances.like(search) |
                                                                                        T_skzi.n_hardware.like(search) | T_skzi.n_destr.like(search) |
                                                                                        T_skzi.note.like(search) | Name_SKZI.n_skzi.like(search)
                                                                                        | T_division.division_name.like(search)
                                                                                        | prof_o.name.like(search)
                                                                                        | prof_i.name.like(search)
                                                                                        | prof_d.name.like(search)
                                                                                        ).all()
    elif q1:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(T_skzi.n_account.like(search_1)).all()
    elif q2:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            Name_SKZI.n_skzi.like(search_2)).all()
    elif q3:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            T_skzi.n_serial.like(search_3)).all()
    elif q4:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            T_skzi.n_instances.like(search_4)).all()
    elif q5:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            T_division.division_name.like(search_5)).all()
    elif q6:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            T_skzi.sopr_pis.like(search_6)).all()
    elif q7:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            prof_o.name.like(search_7)).all()
    elif (q8_1 or q8_2): # --- Поиск по дате ????
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id)

        if ( q8_1 and q8_2): # поля дат не пустые
            if q8_1 <= q8_2:
                sp = sp.filter((T_skzi.date_rec.between(q8_1, q8_2)))
                sp = sp.all()
            else:
                return render_template("Err_msg.html", err=err)
        elif (q8_1 or q8_2):
            sp = sp.filter((T_skzi.date_rec <= (q8_1 + q8_2)))
            sp = sp.all()

            # return render_template("Err_msg.html", err=err)

    elif q9: #
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            prof_i.name.like(search_9)).all()
    # elif q10: #
    #     sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
    #         T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
    #         T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
    #         T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
    #         T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
    #         prof_i.name.like(search_10)).all()
    elif (q10_1 or q10_2): # --- Поиск по дате ????
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id)

        if (q10_1 and q10_2):  # поля дат не пустые
            if q8_1 <= q8_2:
                sp = sp.filter((T_skzi.date_con.between(q10_1, q10_2)))
                sp = sp.all()
            else:
                return render_template("Err_msg.html", err=err)
        elif (q10_1 or q10_2):
            sp = sp.filter((T_skzi.date_con <= (q10_1+q10_2)))
            sp = sp.all()
            # return render_template("Err_msg.html", err=err)
    elif q11: #
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            T_skzi.n_hardware.like(search_11)).all()
    elif q12: #
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            prof_d.name.like(search_12)).all()
    elif (q12_1 or q12_2):  # --- Поиск по дате ????
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id)

        if (q12_1 and q12_2):  # поля дат не пустые
            if q12_1 <= q12_2:
                sp = sp.filter((T_skzi.data_destr.between(q12_1, q12_2)))
                sp = sp.all()
            else:
                return render_template("Err_msg.html", err=err)
        elif (q12_1 or q12_2):
            sp = sp.filter((T_skzi.data_destr <= (q12_1+q12_2)))
            sp = sp.all()
            # return render_template("Err_msg.html", err=err)
    elif q13: #
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            prof_d.name.like(search_13)).all()
    elif q14: #
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            T_skzi.n_destr.like(search_14)).all()
    elif q15: #
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(
            T_skzi.note.like(search_15)).all()

    else:
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).order_by(T_skzi.id.asc()).all()

    # return render_template("list_skzi.html", sp=sp)
    # записываем значение фильтров в куки
    # res = make_response(render_template("list_sk.html", sp=sp, str_id=str_id))
    # if q1: # если q1 не пустойЮ то запишем куки
    #     res.set_cookie('q1', q1)

    # res.set_cookie('q1', '111', max_age=60 * 60 * 24 * 365 * 2)
    # res.set_cookie('q2', '555')
    # return res
    return render_template("list_sk.html", sp=sp, str_id=str_id)



@app.route('/key/<int:id>')
@login_required
def key(id):
    spisok_key = db.session.query(T_Key,  T_license, T_division,TS_CENTER).filter(T_Key.ID_LICENSE==T_license.id).filter(T_Key.KEY_OWNER_ID==T_division.id).filter(T_Key.ID_CENTER==TS_CENTER.id).filter(T_license.id==id).order_by(TS_CENTER.CERTIFICATION_CENTER.desc()).all()
    id_key = id
    lic_f = T_license.query.get(id)
    return render_template("key.html", spisok_key=spisok_key, id_key=id_key, lic_f=lic_f)


@app.route('/adv-search', methods=['POST', 'GET'])
@login_required
def adv_search():

    fin_sh = 'l}qO8qqs'
    err = 'Неверно указан период дат!'
    us = current_user.id

    users_o = aliased(Users)
    users_i = aliased(Users)
    users_d = aliased(Users)
    prof_o = aliased(Profiles)
    prof_i = aliased(Profiles)
    prof_d = aliased(Profiles)

    opt_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi.asc()).all()
    option = T_division.query.order_by(T_division.division_name.asc()).all()
    # post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
    # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    fio = db.session.query(Users, Profiles).filter(Users.id==Profiles.user_id).order_by(Profiles.name.asc()).all()

    div_id = Profiles.query.filter(Profiles.user_id == us).first()

    if request.method == "POST":
        n_serial = request.form['n_skzi']  # - Серийные номера СКЗИ - 4
        n_instances = request.form['n_instances']  # Номера экземпляров (копий) - 5
        sopr_pis = request.form['sopr_pis']  # Дата и номер сопроводительного письма - 7

        date_rec = request.form['date_rec']  # Дата получения - 9
        date_rec2 = request.form['date_rec2']  # Дата получения

        date_con = request.form['date_con']  # Дата подключения (установки) -11
        date_con2 = request.form['date_con2']  # Дата подключения (установки)

        n_hardware = request.form['n_hardware']  # Инвентарный (серийный) номер аппаратных средств -12

        data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
        data_destr2 = request.form['data_destr2']  # Дата изъятия (уничтожения)- 13

        n_destr = request.form['n_destr']  # - Номер акта или расписка об уничтожении - 15
        note = request.form['note']  # - Примечание - 16

        name_skzi = request.form['name_cript']  # key for - Наименование криптосредства - 3
        div_sours = request.form['div_sours']  # key for От кого получены - 6
        fio_owner = request.form['fio_owner']  # key for Ф.И.О. владельца криптосредства - 8
        fio_instal = request.form['fio_instal']  # key for Ф.И.О. производившего подключение (установку) -10
        fio_destr = request.form[
            'fio_destr']  # key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14


        # if n_serial:
        #     n_serial = "%{}%".format(n_serial)
        # else:
        #     n_serial = fin_sh
        #
        # if n_instances:
        #     # n_serial = "%{}%".format(n_serial)
        #     n_instances = "%{}%".format(n_instances)
        # else:
        #     n_instances = fin_sh
        #
        # if n_hardware:
        #     # n_serial = "%{}%".format(n_serial)
        #     n_hardware = "%{}%".format(n_hardware)
        # else:
        #     n_hardware = fin_sh
        #
        # if n_destr:
        #     # n_serial = "%{}%".format(n_serial)
        #     n_destr = "%{}%".format(n_destr)
        # else:
        #     n_destr = fin_sh
        #
        # if note:
        #     n_serial = "%{}%".format(n_serial)
        #     note = "%{}%".format(note)
        # else:
        #     note = fin_sh

        # Если текстовые поля карточки расш. поиска пустые и
        # и поля из справочников не выбраны

        # if name_skzi == div_sours == fio_owner == fio_instal == fio_destr == 1:
        #     if n_serial == n_instances == n_hardware == n_destr == note == fin_sh:
        #         n_serial = n_instances = n_hardware = n_destr = note = "%%"

        # n_serial = "%{}%".format(n_serial)
        # n_instances = "%{}%".format(n_instances)
        # n_hardware = "%{}%".format(n_hardware)
        # n_destr = "%{}%".format(n_destr)
        # note = "%{}%".format(note)

        #--- Если значение полей спр-ков == 1 то по ним искать не надо
        # значение 1 поля будем считать не заданным


        # sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
        #     T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
        #     T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
        #     T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
        #     T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter((
        #     T_skzi.n_serial.like(n_serial) | T_skzi.n_instances.like(n_instances) |
        #     T_skzi.n_hardware.like(n_hardware) | T_skzi.n_destr.like(n_destr) |
        #     T_skzi.note.like(note)))
        l = []
        sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
            T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
            T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
            T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
            T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id)

        if name_skzi > '1':
            # sp = sp.filter(and_(T_skzi.name_skzi== name_skzi))
            sp = sp.filter(T_skzi.name_skzi == name_skzi)

        if fio_owner > '1':
            sp = sp.filter((T_skzi.fio_owner == fio_owner))
        #
        if fio_instal > '1':
                sp = sp.filter(and_(T_skzi.fio_instal == fio_instal))

        if div_sours > '1':
            sp = sp.filter(and_(T_skzi.div_sours == div_sours))
        #
        if fio_destr > '1':
            sp = sp.filter(and_(T_skzi.fio_destr == fio_destr))

        # if (date_rec and date_rec2) or (date_rec <= date_rec2):
        if (date_rec and date_rec2): # поля дат не пустые
            if date_rec <= date_rec2:
                sp = sp.filter((T_skzi.date_rec.between(date_rec, date_rec2)))
            else:
                return render_template("Err_msg.html", err=err)
        elif (date_rec or date_rec2):
            return render_template("Err_msg.html", err=err)

        if (date_con and date_con2): # поля дат не пустые
            if date_con <= date_con2:
                sp = sp.filter((T_skzi.date_con.between(date_con, date_con2)))
            else:
                return render_template("Err_msg.html", err=err)
        elif (date_con or date_con2):
            return render_template("Err_msg.html", err=err)

        if (data_destr and data_destr2): # поля дат не пустые
            if data_destr <= data_destr2:
                sp = sp.filter((T_skzi.data_destr.between(data_destr, data_destr2)))
            else:
                return render_template("Err_msg.html", err=err)
        elif (data_destr or data_destr2):
            return render_template("Err_msg.html", err=err)

        # --- ИЛИ
        if n_serial:
            n_serial = "%{}%".format(n_serial)
            l.append(T_skzi.n_serial.like(n_serial))
            # sp = sp.filter(or_(T_skzi.n_serial.like(n_serial)))

        if n_instances:
            n_instances = "%{}%".format(n_instances)
            # sp = sp.filter(or_(T_skzi.n_instances.like(n_instances)))
            l.append(T_skzi.n_instances.like(n_instances))

        if n_hardware:
            n_hardware = "%{}%".format(n_hardware)
            l.append(T_skzi.n_hardware.like(n_hardware))

        if n_destr:
            n_destr = "%{}%".format(n_destr)
            l.append(T_skzi.n_destr.like(n_destr))

        if note:
            note = "%{}%".format(note)
            l.append(T_skzi.note.like(note))

        sp = sp.filter(or_(*l))

        sp = sp.all()

        # s_count2 = len(sp)

            # | ((T_skzi.name_skzi == name_skzi))).all()
            # & (T_skzi.fio_owner == fio_owner ))).all()
            # | T_division.division_name.like(search)
            # | prof_o.name.like(search)
            # | prof_i.name.like(search)
            # | prof_d.name.like(search)
            # ).all()

        return render_template("search_res.html", sp=sp)
    else:
        return render_template("adv_search.html", opt_skzi=opt_skzi, option=option, fio=fio)


# @app.route('/test-add', methods=['POST', 'GET'])
# @login_required
# def test_add():
@app.route('/create-skzi', methods=['POST', 'GET'])
@login_required
def create_skzi():
    global selStr
    error = None
    err = [0,0,0,0,0,0,0,0,0,0,0,0,0,0] # это для ошибок
    us = current_user.id
    opt_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi.asc()).all()
    option = T_division.query.order_by(T_division.division_name.asc()).all()
    # post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
    # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    fio = db.session.query(Users, Profiles).filter(Users.id==Profiles.user_id).order_by(Profiles.name.asc()).all()

    div_id = Profiles.query.filter(Profiles.user_id == us).first()

    if request.method == "POST":

        # if len(request.form['n_skzi']) > 1 and len(request.form['n_instances']) > 1 and len(
        #         request.form['sopr_pis']) > 1 \
        #         and len(request.form['n_instances']) > 1 and len(request.form['n_instances']) > 1\
        #         and len(request.form['n_hardware']) > 1  and len(request.form['n_destr']) > 1\
        #         and len(request.form['note']) > 1\
        #         and validate(request.form['date_rec']) \
        #         and validate(request.form['date_con'])\
        #         and validate(request.form['data_destr']):

            # n_account =   # - учетный номер - 2
        n_serial =  request.form['n_skzi'] # - Серийные номера СКЗИ - 4
        n_instances =  request.form['n_instances'] # Номера экземпляров (копий) - 5
        sopr_pis = request.form['sopr_pis'] # Дата и номер сопроводительного письма - 7
        date_rec = request.form['date_rec'] # Дата получения - 9
        date_con = request.form['date_con'] # Дата подключения (установки) -11
        n_hardware = request.form['n_hardware'] # Инвентарный (серийный) номер аппаратных средств -12
        data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
        n_destr = request.form['n_destr'] # - Номер акта или расписка об уничтожении - 15
        note = request.form['note'] # - Примечание - 16

        name_skzi =  request.form['name_cript']# key for - Наименование криптосредства - 3
        div_sours = request.form['div_sours'] # key for От кого получены - 6
        fio_owner =  request.form['fio_owner']# key for Ф.И.О. владельца криптосредства - 8
        fio_instal = request.form['fio_instal'] # key for Ф.И.О. производившего подключение (установку) -10
        fio_destr =  request.form['fio_destr']# key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14

            # print(request.form.getlist('list'))

        if not name_skzi or (name_skzi == '1'):
                error = 'Наименование криптосредства - не выбрано'
                flash(error, category="error")
                err[0] = 1
        if not n_serial or not n_serial.strip():
                error = 'Номер СКЗи - отсутствует'
                flash(error, category="error")
                err[1] = 1
        if not n_instances or not n_instances.strip():
                error = 'Номера экземпляров - отсутствует'
                flash(error, category="error")
                err[2] = 1
        if not div_sours or (div_sours == '1'):
            error = 'От кого получено - не выбрано'
            flash(error, category="error")
            err[3] = 1
        if not sopr_pis or not sopr_pis.strip():
                error = 'Дата и номер сопроводительного письма - отсутствует'
                flash(error, category="error")
                err[4] = 1
        if not fio_owner or (fio_owner == '1'):
            error = 'Ф.И.О. владельца криптосредства - не выбрано'
            flash(error, category="error")
            err[5] = 1
        if not fio_instal or (fio_instal == '1'):
            error = 'Ф.И.О. производившего подключение (установку) - не выбрано'
            flash(error, category="error")
            err[6] = 1
        if not validate(date_rec):
                error = 'Дата получения - отсутствует'
                flash(error, category="error")
                err[7] = 1
        if not n_hardware or not n_hardware.strip():
                error = 'Инвентарный (серийный) номер аппаратных средств - отсутствует'
                flash(error, category="error")
                err[8] = 1
        if not validate(date_con):
                error = 'Дата подключения (установки)- отсутствует'
                flash(error, category="error")
                err[9] = 1
        if not fio_destr or (fio_destr == '1'):
                error = 'Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - отсутствует'
                flash(error, category="error")
                err[10] = 1
        if not validate(data_destr):
                error = ' Дата изъятия (уничтожения)- отсутствует'
                flash(error, category="error")
                err[11] = 1
        if not  n_destr or not  n_destr.strip():
                error = 'Номер акта или расписка об уничтожении - отсутствует'
                flash(error, category="error")
                err[12] = 1
        if not note or not note.strip():
                error = 'Примечание - отсутствует'
                flash(error, category="error")
                err[13] = 1

        if not error:
            # Это для нового номера СКЗИ
            val_uuid = '\''+str(div_id.DIVISION_OWNER)+'\''
            query = 'select gen_number('+val_uuid+')'
            results = db.session.execute(query)
            for r in results:
                n_account = r[0]

            sp_skz = T_skzi(n_serial=n_serial, n_account=n_account, n_instances=n_instances,  sopr_pis=sopr_pis, date_rec=date_rec, date_con=date_con,
            n_hardware=n_hardware, data_destr=data_destr, n_destr=n_destr, note=note, name_skzi=int(name_skzi), div_sours=int(div_sours), fio_owner=int(fio_owner), fio_instal=int(fio_instal), fio_destr=int(fio_destr))

            # sp = request.form.getlist('list')

            us = current_user.id

            try:
                # db.session.add(sp_skz)
                # db.session.flush()
                #
                # db.session.refresh(sp_skz)
                # # print(lic.id)


                # for s in request.form.getlist('list'):
                #     print(s)
                #     sys = TC_LICENSE_SYSTEM(SYSTEM_ID=s, LICENSE_ID=lic.id)
                #     db.session.add(sys)

                db.session.add(sp_skz)
                db.session.flush()

                log = T_log(LICENSE_ID=sp_skz.id, EVENT_TYPE=1, COMMENT='', USER_ID=us)
                db.session.add(log)

                db.session.commit()

                # if par == 'Buy':
                #     return redirect('/license/' + str(lic.id) + '/update')
                # else:
                selStr = sp_skz.id  # Для передачи в обработчик List
                return redirect('/list')

            except:
                return "При добавление записи произошла ошибка"

        else: #                                                                 ---   Переход если не заполнены поля

             # context = request.form
             # s = context['name_cript']
             # l_err = [0,1]
             return render_template("test_add.html", opt_skzi=opt_skzi, option=option, fio=fio, context = request.form, err=err)
            # return "Неверно заполнены поля"

    else:
        # l_err = [0, 1]
        # context = request.form
        return render_template("test_add.html", opt_skzi=opt_skzi, option=option, fio=fio,  context = request.form, err=err)


#" Тест для update с валидацией
# @app.route('/test-update/<int:id>/update', methods=['POST', 'GET'])
# @login_required
# def test_update(id):
@app.route('/list/<int:id>/update', methods=['POST', 'GET'])
@login_required
def list_update(id):
    global selStr
    error = None
    err = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]  # это для ошибок
    us = current_user.id

    users_o = aliased(Users)
    users_i = aliased(Users)
    users_d = aliased(Users)
    prof_o = aliased(Profiles)
    prof_i = aliased(Profiles)
    prof_d = aliased(Profiles)

    sp = db.session.query(T_skzi, Name_SKZI, T_division, users_o, users_i, users_d, prof_o, prof_i, prof_d).filter(
        T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
        T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
        T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
        T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(T_skzi.id == id).first()

    opt_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi.asc()).all()
    option = T_division.query.order_by(T_division.division_name.asc()).all()
    # post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
    # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    fio = db.session.query(Users, Profiles).filter(Users.id == Profiles.user_id).order_by(Profiles.name.asc()).all()

    sp2 = T_skzi.query.get(id)

    if request.method == "POST":
        n_serial = request.form['n_skzi']  # - Серийные номера СКЗИ - 4
        n_instances = request.form['n_instances']  # Номера экземпляров (копий) - 5
        sopr_pis = request.form['sopr_pis']  # Дата и номер сопроводительного письма - 7
        date_rec = request.form['date_rec']  # Дата получения - 9
        date_con = request.form['date_con']  # Дата подключения (установки) -11
        n_hardware = request.form['n_hardware']  # Инвентарный (серийный) номер аппаратных средств -12
        data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
        n_destr = request.form['n_destr']  # - Номер акта или расписка об уничтожении - 15
        note = request.form['note']  # - Примечание - 16

        name_skzi = request.form['name_cript']  # key for - Наименование криптосредства - 3
        div_sours = request.form['div_sours']  # key for От кого получены - 6
        fio_owner = request.form['fio_owner']  # key for Ф.И.О. владельца криптосредства - 8
        fio_instal = request.form['fio_instal']  # key for Ф.И.О. производившего подключение (установку) -10
        fio_destr = request.form['fio_destr']  # key

        if not name_skzi or (name_skzi == '1'):
            error = 'Наименование криптосредства - не выбрано'
            flash(error, category="error")
            err[0] = 1
        if not n_serial or not n_serial.strip():
            error = 'Номер СКЗи - отсутствует'
            flash(error, category="error")
            err[1] = 1
        if not n_instances or not n_instances.strip():
            error = 'Номера экземпляров - отсутствует'
            flash(error, category="error")
            err[2] = 1
        if not div_sours or (div_sours == '1'):
            error = 'От кого получено - не выбрано'
            flash(error, category="error")
            err[3] = 1
        if not sopr_pis or not sopr_pis.strip():
            error = 'Дата и номер сопроводительного письма - отсутствует'
            flash(error, category="error")
            err[4] = 1
        if not fio_owner or (fio_owner == '1'):
            error = 'Ф.И.О. владельца криптосредства - не выбрано'
            flash(error, category="error")
            err[5] = 1
        if not fio_instal or (fio_instal == '1'):
            error = 'Ф.И.О. производившего подключение (установку) - не выбрано'
            flash(error, category="error")
            err[6] = 1
        if not validate(date_rec):
            error = 'Дата получения - отсутствует'
            flash(error, category="error")
            err[7] = 1
        if not n_hardware or not n_hardware.strip():
            error = 'Инвентарный (серийный) номер аппаратных средств - отсутствует'
            flash(error, category="error")
            err[8] = 1
        if not validate(date_con):
            error = 'Дата подключения (установки)- отсутствует'
            flash(error, category="error")
            err[9] = 1
        if not fio_destr or (fio_destr == '1'):
            error = 'Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - отсутствует'
            flash(error, category="error")
            err[10] = 1
        if not validate(data_destr):
            error = ' Дата изъятия (уничтожения)- отсутствует'
            flash(error, category="error")
            err[11] = 1
        if not n_destr or not n_destr.strip():
            error = 'Номер акта или расписка об уничтожении - отсутствует'
            flash(error, category="error")
            err[12] = 1
        if not note or not note.strip():
            error = 'Примечание - отсутствует'
            flash(error, category="error")
            err[13] = 1
        elif 0 in err: # если нет ошибок
        # if len(request.form['n_skzi']) > 1 and len(request.form['n_instances']) > 1 and len(
        #         request.form['sopr_pis']) > 1 \
        #         and len(request.form['n_instances']) > 1 and len(request.form['n_instances']) > 1 \
        #         and len(request.form['n_hardware']) > 1 and len(request.form['n_destr']) > 1 \
        #         and len(request.form['note']) > 1 \
        #         and validate(request.form['date_rec']) \
        #         and validate(request.form['date_con']) \
        #         and validate(request.form['data_destr']):
        #     n_account =   # - учетный номер - 2
        #     sp2.n_serial = request.form['n_skzi']  # - Серийные номера СКЗИ - 4
            sp2.n_instances = request.form['n_instances']  # Номера экземпляров (копий) - 5
            sp2.sopr_pis = request.form['sopr_pis']  # Дата и номер сопроводительного письма - 7
            sp2.date_rec = request.form['date_rec']  # Дата получения - 9
            sp2.date_con = request.form['date_con']  # Дата подключения (установки) -11
            sp2.n_hardware = request.form['n_hardware']  # Инвентарный (серийный) номер аппаратных средств -12
            sp2.data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
            sp2.n_destr = request.form['n_destr']  # - Номер акта или расписка об уничтожении - 15
            sp2.note = request.form['note']  # - Примечание - 16

            sp2.name_skzi = request.form['name_cript']  # key for - Наименование криптосредства - 3
            sp2.div_sours = request.form['div_sours']  # key for От кого получены - 6
            sp2.fio_owner = request.form['fio_owner']  # key for Ф.И.О. владельца криптосредства - 8
            sp2.fio_instal = request.form['fio_instal']  # key for Ф.И.О. производившего подключение (установку) -10
            sp2.fio_destr = request.form[
                'fio_destr']  # key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14

            try:
                log = T_log(LICENSE_ID=id, EVENT_TYPE=3, COMMENT='', USER_ID=us)
                db.session.add(log)
                db.session.commit()
                selStr = id
                return redirect('/list')
            except:
                return "При редактировании записи произошла ошибка"
        else:
            return render_template("test_update.html", sp=sp, option=option, fio=fio, opt_skzi=opt_skzi, context=request.form, err=err)
            # flash('Неверно заполнены поля', category="error")
            # return render_template("list_update.html", sp=sp, option=option, fio=fio, opt_skzi=opt_skzi)
            # return render_template("test_update.html", context=request.form)
    else:#                                                                 ---   Переход если не заполнены поля ---
        return render_template("test_update.html", sp=sp, option=option, fio=fio, opt_skzi=opt_skzi, context = request.form, err=err)


# @app.route('/create-skzi', methods=['POST', 'GET'])
# @login_required
# def create_skzi():
#     global selStr
#
#     us = current_user.id
#     opt_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi.asc()).all()
#     option = T_division.query.order_by(T_division.division_name.asc()).all()
#     # post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
#     # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
#     fio = db.session.query(Users, Profiles).filter(Users.id==Profiles.user_id).order_by(Profiles.name.asc()).all()
#
#     div_id = Profiles.query.filter(Profiles.user_id == us).first()
#
#     if request.method == "POST":
#
#         if len(request.form['n_skzi']) > 1 and len(request.form['n_instances']) > 1 and len(
#                 request.form['sopr_pis']) > 1 \
#                 and len(request.form['n_instances']) > 1 and len(request.form['n_instances']) > 1\
#                 and len(request.form['n_hardware']) > 1  and len(request.form['n_destr']) > 1\
#                 and len(request.form['note']) > 1\
#                 and validate(request.form['date_rec']) \
#                 and validate(request.form['date_con'])\
#                 and validate(request.form['data_destr']):
#
#             # n_account =   # - учетный номер - 2
#             n_serial =  request.form['n_skzi'] # - Серийные номера СКЗИ - 4
#             n_instances =  request.form['n_instances'] # Номера экземпляров (копий) - 5
#             sopr_pis = request.form['sopr_pis'] # Дата и номер сопроводительного письма - 7
#             date_rec = request.form['date_rec'] # Дата получения - 9
#             date_con = request.form['date_con'] # Дата подключения (установки) -11
#             n_hardware = request.form['n_hardware'] # Инвентарный (серийный) номер аппаратных средств -12
#             data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
#             n_destr = request.form['n_destr'] # - Номер акта или расписка об уничтожении - 15
#             note = request.form['note'] # - Примечание - 16
#
#             name_skzi =  request.form['name_cript']# key for - Наименование криптосредства - 3
#             div_sours = request.form['div_sours'] # key for От кого получены - 6
#             fio_owner =  request.form['fio_owner']# key for Ф.И.О. владельца криптосредства - 8
#             fio_instal = request.form['fio_instal'] # key for Ф.И.О. производившего подключение (установку) -10
#             fio_destr =  request.form['fio_destr']# key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14
#
#             # print(request.form.getlist('list'))
#
#
#             val_uuid = '\''+str(div_id.DIVISION_OWNER)+'\''
#             query = 'select gen_number('+val_uuid+')'
#             results = db.session.execute(query)
#             for r in results:
#                 n_account = r[0]
#
#             sp_skz = T_skzi(n_serial=n_serial, n_account=n_account, n_instances=n_instances,  sopr_pis=sopr_pis, date_rec=date_rec, date_con=date_con,
#             n_hardware=n_hardware, data_destr=data_destr, n_destr=n_destr, note=note, name_skzi=int(name_skzi), div_sours=int(div_sours), fio_owner=int(fio_owner), fio_instal=int(fio_instal), fio_destr=int(fio_destr))
#
#             # sp = request.form.getlist('list')
#
#             us = current_user.id
#
#             try:
#                 # db.session.add(sp_skz)
#                 # db.session.flush()
#                 #
#                 # db.session.refresh(sp_skz)
#                 # # print(lic.id)
#
#
#                 # for s in request.form.getlist('list'):
#                 #     print(s)
#                 #     sys = TC_LICENSE_SYSTEM(SYSTEM_ID=s, LICENSE_ID=lic.id)
#                 #     db.session.add(sys)
#
#                 db.session.add(sp_skz)
#                 db.session.flush()
#
#                 log = T_log(LICENSE_ID=sp_skz.id, EVENT_TYPE=1, COMMENT='', USER_ID=us)
#                 db.session.add(log)
#
#                 db.session.commit()
#
#                 # if par == 'Buy':
#                 #     return redirect('/license/' + str(lic.id) + '/update')
#                 # else:
#                 selStr = sp_skz.id # Для передачи в обработчик List
#                 return redirect('/list')
#
#             except:
#                 return "При добавление записи произошла ошибка"
#
#         else:
#              flash('Неверно заполнены поля', category="error")
#              return render_template("create-skzi.html", opt_skzi=opt_skzi, option=option, fio=fio)
#             # return "Неверно заполнены поля"
#
#     else:
#         return render_template("create-skzi.html", opt_skzi=opt_skzi, option=option, fio=fio)


@app.route('/journal')
@login_required
def journal():
    # text = request.args.get('button_text')
    # print()
    # print('Button text:', text)
    # print()
    # spisok_key = db.session.query(T_Key,  T_license, T_division,TS_CENTER).filter(T_Key.ID_LICENSE==T_license.id).filter(T_Key.KEY_OWNER_ID==T_division.id).filter(T_Key.ID_CENTER==TS_CENTER.id).filter(T_license.id==id).order_by(TS_CENTER.CERTIFICATION_CENTER.desc()).all()
    # id_key = id
    # lic_f = T_license.query.get(id)
    q = request.args.get('q')
    search = "%{}%".format(q)

    a = aliased(T_division)

    if q:
        res = db.session.query(Journal_skzi, T_division, a, T_license).join(T_division,
                                                                        Journal_skzi.DIVISION_ID == T_division.id).join(
        a, Journal_skzi.DIV_SOUR_ID == a.id).join(T_license, Journal_skzi.LICENSE_ID == T_license.id).filter(T_license.code.like(search) | T_division.division_name.like(search) | a.division_name.like(search) ).order_by(
        Journal_skzi.data_begin.asc(), Journal_skzi.uuid.asc()).all()
    else:
        res = db.session.query(Journal_skzi, T_division, a, T_license).join(T_division,
                                                                        Journal_skzi.DIVISION_ID == T_division.id).join(
        a, Journal_skzi.DIV_SOUR_ID == a.id).join(T_license, Journal_skzi.LICENSE_ID == T_license.id).order_by(
        Journal_skzi.data_begin.asc(), Journal_skzi.uuid.asc()).all()

    return render_template("journal.html", res=res)
    # return render_template("journal.html", spisok_key=spisok_key, id_key=id_key, lic_f=lic_f)


@app.route('/list/<int:id>')
@login_required
def list_detail(id):

    us = current_user.id

    users_o = aliased(Users)
    users_i = aliased(Users)
    users_d = aliased(Users)
    prof_o = aliased(Profiles)
    prof_i = aliased(Profiles)
    prof_d = aliased(Profiles)

    sp = db.session.query(T_skzi, Name_SKZI, T_division, users_o, users_i, users_d, prof_o, prof_i, prof_d).filter(
        T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
        T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
        T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
        T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(T_skzi.id == id).first()

    opt_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi.asc()).all()
    option = T_division.query.order_by(T_division.division_name.asc()).all()
    # post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
    # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    fio = db.session.query(Users, Profiles).filter(Users.id == Profiles.user_id).order_by(Profiles.name.asc()).all()

    # div_id = Profiles.query.filter(Profiles.user_id == us).first()
    # sp2 = T_skzi.query.get(id)
    return render_template("list_detail.html", sp=sp, option=option, fio=fio, opt_skzi=opt_skzi)


@app.route('/license/<int:id>')
@login_required
def license_detail(id):
    # lic = T_license.query.get(id)
    # lic = db.session.query(T_license, T_division, Post).outerjoin(T_division, T_division.id == T_license.id).outerjoin(
    #     Post, Post.id == T_license.id).filter(T_license.id == id).first()
    a = aliased(T_division)
    lic = db.session.query(T_license, T_division, a, Post).filter(T_division.id == T_license.div_id).filter(
        a.id == T_license.license_owner_id).filter(Post.id == T_license.post_id).filter(T_license.id == id).first()

    option = T_division.query.order_by(T_division.division_name.desc()).all()
    post_sp = Post.query.order_by(Post.POST_NAME.desc()).all()
    sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    sys_con = db.session.query(TC_LICENSE_SYSTEM.SYSTEM_ID).filter(TC_LICENSE_SYSTEM.LICENSE_ID == id).all()

    return render_template("lic_detail.html", lic=lic, option=option, post_sp=post_sp, sys_sp=sys_sp, sys_con=sys_con)


@app.route('/posts/<int:id>')
@login_required
def post_detail(id):
    article = Article.query.get(id)
    return render_template("posts_detail.html", article=article)


@app.route('/posts/<int:id>/del')
@login_required
def post_delete(id):
    article = Article.query.get_or_404(id)

    try:
        db.session.delete(article)
        db.session.commit()
        return redirect('/posts')
    except:
        return "При удалении статьи произошла ошибка"


@app.route('/list/<int:id>/del')
@login_required
def list_delete(id):
    us = current_user.id
    licen = T_skzi.query.get_or_404(id)

    # sql1 = delete(TC_LICENSE_SYSTEM).where(TC_LICENSE_SYSTEM.LICENSE_ID == id)
    # sql2 = delete(T_Key).where(T_Key.ID_LICENSE == id)
    log = T_log(LICENSE_ID=id, EVENT_TYPE=4, COMMENT='', USER_ID=us)

    av = Journal_skzi.query.filter(Journal_skzi.LICENSE_ID==id).first()
    # if av:
    #     return ("Удалить лицензию нельзя! Она в журнале СКЗИ")

    try:
        # db.session.execute(sql1)
        # db.session.execute(sql2)
        db.session.delete(licen)
        db.session.add(log)
        db.session.commit()
        return redirect('/list')
    except:
        return "При удалении записи произошла ошибка"


@app.route('/license/<int:id>/del')
@login_required
def license_delete(id):
    us = current_user.id
    licen = T_license.query.get_or_404(id)
    sql1 = delete(TC_LICENSE_SYSTEM).where(TC_LICENSE_SYSTEM.LICENSE_ID == id)
    sql2 = delete(T_Key).where(T_Key.ID_LICENSE == id)
    log = T_log(LICENSE_ID=id, EVENT_TYPE=4, COMMENT='', USER_ID=us)

    av = Journal_skzi.query.filter(Journal_skzi.LICENSE_ID==id).first()

    if av:
        return("Удалить лицензию нельзя! Она в журнале СКЗИ")

    try:
        db.session.execute(sql1)
        db.session.execute(sql2)
        db.session.delete(licen)
        db.session.add(log)
        db.session.commit()
        return redirect('/license')
    except:
        return "При удалении лицензиии произошла ошибка"


@app.route('/journal/<uuid>/del')
@login_required
def journal_delete(uuid):
    us = current_user.id
    # journ = Journal_skzi.query.get_or_404(id)
    journ = db.session.query(Journal_skzi).filter(Journal_skzi.uuid == uuid).all()
    # qwe = journ.LICENSE_ID
    # log = T_log(LICENSE_ID=journ.LICENSE_ID, EVENT_TYPE=6, COMMENT='', USER_ID=us)
    sql1 = delete(Journal_skzi).where(Journal_skzi.uuid==uuid)

    #
    # for r in journ:
    #     w = r.LICENSE_ID
    #     log = T_log(LICENSE_ID=r.LICENSE_ID, EVENT_TYPE=6, COMMENT='', USER_ID=us)
    #     db.session.add(log)
    #
    # db.session.execute(sql1)
    # db.session.commit()
    #
    # return "При удалении документа произошла ошибка"

    try:
        for r in journ:
            w = r.LICENSE_ID
            log = T_log(LICENSE_ID=r.LICENSE_ID, EVENT_TYPE=6, COMMENT='', USER_ID=us)
            db.session.add(log)

        # db.session.delete(journ)
        db.session.execute(sql1)
        db.session.commit()
        return redirect('/journal')
    except:
        return "При удалении документа произошла ошибка"


# @app.route('/view/<var1>/<var2>')
# def view(var1, var2):
@app.route('/view/')
def view():
    pass
    return redirect('/license/')


# http://127.0.0.1:5000/key/6/del
@app.route('/key/<int:id_k>/<int:id_l>/del')
@login_required
def key_delete(id_k, id_l):
    us = current_user.id
    key_del = T_Key.query.get_or_404(id_k)
    log = T_log(LICENSE_ID=id_l, EVENT_TYPE=5, COMMENT='', USER_ID=us)

    try:
        db.session.delete(key_del)
        db.session.add(log)

        db.session.commit()
        # return redirect('/key/26' + str(id))
        return redirect('/key/' + str(id_l))
    except:
        return "При удалении ключа произошла ошибка"


@app.route('/posts/<int:id>/update', methods=['POST', 'GET'])
@login_required
def post_update(id):
    article = Article.query.get(id)
    if request.method == "POST":
        article.title = request.form['title']
        article.intro = request.form['intro']
        article.text = request.form['text']

        try:
            db.session.commit()
            return redirect('/posts')
        except:
            return "При редактировании статьи произошла ошибка"
    else:

        return render_template("post_update.html", article=article)


@app.route('/list/<int:id>/copy', methods=['POST', 'GET'])
@login_required
def lic_copy(id):
    global selStr

    us = current_user.id

    users_o = aliased(Users)
    users_i = aliased(Users)
    users_d = aliased(Users)
    prof_o = aliased(Profiles)
    prof_i = aliased(Profiles)
    prof_d = aliased(Profiles)

    sp = db.session.query(T_skzi, Name_SKZI, T_division, users_o, users_i, users_d, prof_o, prof_i, prof_d).filter(
        T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
        T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
        T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
        T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(T_skzi.id == id).first()

    opt_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi.asc()).all()
    option = T_division.query.order_by(T_division.division_name.asc()).all()
    # post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
    # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    fio = db.session.query(Users, Profiles).filter(Users.id == Profiles.user_id).order_by(Profiles.name.asc()).all()

    div_id = Profiles.query.filter(Profiles.user_id == us).first()

    # sp2 = T_skzi.query.get(id)

    if request.method == "POST":

        # fio = request.form['fio'] # ФИО
        # div_id = request.form['division'] # Подразделение
        # post_id = request.form['post'] # Должность
        # license_code = request.form['lic_code'] # Лицензия
        # license_owner_id = request.form['lic_owner_id'] # Где оригинал
        # computer_name = request.form['name_comp'] # Имя компьютера
        # comment = request.form['comment'] # Примечание
        # contract = request.form['contract'] # Контракт
        # code = request.form['code'] # Код
        #
        # par = request.form['submitType']

        # n_account =   # - учетный номер - 2


        # sp2.n_serial =  request.form['n_skzi'] # - Серийные номера СКЗИ - 4
        # sp2.n_instances =  request.form['n_instances'] # Номера экземпляров (копий) - 5
        # sp2.sopr_pis = request.form['sopr_pis'] # Дата и номер сопроводительного письма - 7
        # sp2.date_rec = request.form['date_rec'] # Дата получения - 9
        # sp2.date_con = request.form['date_con'] # Дата подключения (установки) -11
        # sp2.n_hardware = request.form['n_hardware'] # Инвентарный (серийный) номер аппаратных средств -12
        # sp2.data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
        # sp2.n_destr = request.form['n_destr'] # - Номер акта или расписка об уничтожении - 15
        # sp2.note = request.form['note'] # - Примечание - 16
        #
        # sp2.name_skzi =  request.form['name_cript']# key for - Наименование криптосредства - 3
        # sp2.div_sours = request.form['div_sours'] # key for От кого получены - 6
        # sp2.fio_owner =  request.form['fio_owner']# key for Ф.И.О. владельца криптосредства - 8
        # sp2.fio_instal = request.form['fio_instal'] # key for Ф.И.О. производившего подключение (установку) -10
        # sp2.fio_destr =  request.form['fio_destr']# key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14

        n_serial = request.form['n_skzi']  # - Серийные номера СКЗИ - 4
        n_instances = request.form['n_instances']  # Номера экземпляров (копий) - 5
        sopr_pis = request.form['sopr_pis']  # Дата и номер сопроводительного письма - 7
        date_rec = request.form['date_rec']  # Дата получения - 9
        date_con = request.form['date_con']  # Дата подключения (установки) -11
        n_hardware = request.form['n_hardware']  # Инвентарный (серийный) номер аппаратных средств -12
        data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
        n_destr = request.form['n_destr']  # - Номер акта или расписка об уничтожении - 15
        note = request.form['note']  # - Примечание - 16

        name_skzi = request.form['name_cript']  # key for - Наименование криптосредства - 3
        div_sours = request.form['div_sours']  # key for От кого получены - 6
        fio_owner = request.form['fio_owner']  # key for Ф.И.О. владельца криптосредства - 8
        fio_instal = request.form['fio_instal']  # key for Ф.И.О. производившего подключение (установку) -10
        fio_destr = request.form[
            'fio_destr']  # key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14

        # print(request.form.getlist('list'))

        val_uuid = '\''+str(div_id.DIVISION_OWNER)+'\''
        query = 'select gen_number('+val_uuid+')'
        results = db.session.execute(query)
        for r in results:
            n_account = r[0]

        sp_skz = T_skzi(n_serial=n_serial, n_account=n_account, n_instances=n_instances, sopr_pis=sopr_pis, date_rec=date_rec,
                        date_con=date_con,
                        n_hardware=n_hardware, data_destr=data_destr, n_destr=n_destr, note=note,
                        name_skzi=int(name_skzi), div_sours=int(div_sours), fio_owner=int(fio_owner),
                        fio_instal=int(fio_instal), fio_destr=int(fio_destr))

        # sp = request.form.getlist('list')

         # sp = request.form.getlist('list')

        try:

            # db.session.add(sp_skz)

            db.session.add(sp_skz)
            db.session.flush()

            db.session.refresh(sp_skz)
            # print(lic.id)

            log = T_log(LICENSE_ID=sp_skz.id, EVENT_TYPE=9, COMMENT='', USER_ID=us)

            db.session.add(log)

            db.session.commit()

            selStr=sp_skz.id

            return redirect('/list')
            # if par == 'Buy':
            #     return redirect('/license/' + str(lic.id) + '/update')
            # else:
            #     return redirect('/license')

        except:
            return "При добавление записи произошла ошибка"
    else:
        return render_template("skzi_copy.html", sp=sp, option=option, fio=fio, opt_skzi=opt_skzi)


# @app.route('/list/<int:id>/update', methods=['POST', 'GET'])
# @login_required
# def list_update(id):
#     global selStr
#     us = current_user.id
#
#     users_o = aliased(Users)
#     users_i = aliased(Users)
#     users_d = aliased(Users)
#     prof_o = aliased(Profiles)
#     prof_i = aliased(Profiles)
#     prof_d = aliased(Profiles)
#
#     sp = db.session.query(T_skzi, Name_SKZI, T_division, users_o, users_i, users_d, prof_o, prof_i, prof_d).filter(
#         T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
#         T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
#         T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
#         T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).filter(T_skzi.id == id).first()
#
#     opt_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi.asc()).all()
#     option = T_division.query.order_by(T_division.division_name.asc()).all()
#     # post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
#     # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
#     fio = db.session.query(Users, Profiles).filter(Users.id == Profiles.user_id).order_by(Profiles.name.asc()).all()
#
#     sp2 = T_skzi.query.get(id)
#
#     if request.method == "POST":
#
#         if len(request.form['n_skzi']) > 1 and len(request.form['n_instances']) > 1 and len(
#                 request.form['sopr_pis']) > 1 \
#                 and len(request.form['n_instances']) > 1 and len(request.form['n_instances']) > 1\
#                 and len(request.form['n_hardware']) > 1  and len(request.form['n_destr']) > 1\
#                 and len(request.form['note']) > 1\
#                 and validate(request.form['date_rec']) \
#                 and validate(request.form['date_con'])\
#                 and validate(request.form['data_destr']):
#             # n_account =   # - учетный номер - 2
#             sp2.n_serial = request.form['n_skzi']  # - Серийные номера СКЗИ - 4
#             sp2.n_instances = request.form['n_instances']  # Номера экземпляров (копий) - 5
#             sp2.sopr_pis = request.form['sopr_pis']  # Дата и номер сопроводительного письма - 7
#             sp2.date_rec = request.form['date_rec']  # Дата получения - 9
#             sp2.date_con = request.form['date_con']  # Дата подключения (установки) -11
#             sp2.n_hardware = request.form['n_hardware']  # Инвентарный (серийный) номер аппаратных средств -12
#             sp2.data_destr = request.form['data_destr']  # Дата изъятия (уничтожения)- 13
#             sp2.n_destr = request.form['n_destr']  # - Номер акта или расписка об уничтожении - 15
#             sp2.note = request.form['note']  # - Примечание - 16
#
#             sp2.name_skzi = request.form['name_cript']  # key for - Наименование криптосредства - 3
#             sp2.div_sours = request.form['div_sours']  # key for От кого получены - 6
#             sp2.fio_owner = request.form['fio_owner']  # key for Ф.И.О. владельца криптосредства - 8
#             sp2.fio_instal = request.form['fio_instal']  # key for Ф.И.О. производившего подключение (установку) -10
#             sp2.fio_destr = request.form['fio_destr']  # key for  Ф.И.О. ответственного, производившего изъятие (уничтожение) СКЗИ - 14
#
#             try:
#                 log = T_log(LICENSE_ID=id, EVENT_TYPE=3, COMMENT='', USER_ID=us)
#                 db.session.add(log)
#                 db.session.commit()
#                 selStr = id # присвоение id строки гл. переменной selStr
#                 return redirect('/list')
#             except:
#                 return "При редактировании записи произошла ошибка"
#         else:
#             flash('Неверно заполнены поля', category="error")
#             # return render_template("list_update.html", sp=sp, option=option, fio=fio, opt_skzi=opt_skzi)
#             return render_template("list_update.html", context=request.form)
#     else:
#         return render_template("list_update.html", sp=sp, option=option, fio=fio, opt_skzi=opt_skzi)


@app.route('/license/<int:id>/update', methods=['POST', 'GET'])
@login_required
def lic_update(id):

    a = aliased(T_division)

    lic = db.session.query(T_license, T_division, a, Post).filter(T_division.id == T_license.div_id).filter(
        a.id == T_license.license_owner_id).filter(Post.id == T_license.post_id).filter(T_license.id == id).first()

    lic2 = T_license.query.get(id)

    option = T_division.query.order_by(T_division.division_name.desc()).all()
    post_sp = Post.query.order_by(Post.POST_NAME.desc()).all()
    sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    sys_con = db.session.query(TC_LICENSE_SYSTEM.SYSTEM_ID).filter(TC_LICENSE_SYSTEM.LICENSE_ID == id).all()

    if request.method == "POST":
        # lic.fio = request.form['fio']
        # lic.computer_name = request.form['computer_name']
        # lic.license_code = request.form['license_code']
        # lic.contract = request.form['contract']
        # lic.comment = request.form['comment']
        # lic.code = request.form['code']
        # lic.div_id = request.form['list_name']
        lic2.fio = request.form['fio'] # ФИО
        lic2.div_id = request.form['division'] # Подразделение
        lic2.post_id = request.form['post'] # Должность
        lic2.license_code = request.form['lic_code'] # Лицензия
        lic2.license_owner_id = request.form['lic_owner_id'] # Где оригинал
        lic2.computer_name = request.form['name_comp'] # Имя компьютера
        lic2.comment = request.form['comment'] # Примечание
        lic2.contract = request.form['contract'] # Контракт
        lic2.code = request.form['code'] # Код
        spis1 = request.form.getlist('list')

        us = current_user.id
        log = T_log(LICENSE_ID=id, EVENT_TYPE=3, COMMENT='', USER_ID=us)

        prov = TC_LICENSE_SYSTEM.query.filter(TC_LICENSE_SYSTEM.LICENSE_ID == id).first()

        try:

            # db.session.add(lic2)
            # db.session.flush()
            #
            # db.session.refresh(lic)
            # # print(lic.id)

            if prov:
                sql = delete(TC_LICENSE_SYSTEM).where(TC_LICENSE_SYSTEM.LICENSE_ID == id)
                db.session.execute(sql)

            for s in spis1:
              #  print(s)
                sys = TC_LICENSE_SYSTEM(SYSTEM_ID=s, LICENSE_ID=id)
                db.session.add(sys)

            db.session.add(log)
            db.session.commit()
            return redirect('/license')
        except:
            return "При редактировании записи произошла ошибка"
    else:

        return render_template("lic_update.html", lic=lic, option=option, post_sp=post_sp, sys_sp=sys_sp, sys_con=sys_con)


@app.route('/create-article', methods=['POST', 'GET'])
@login_required
def create_article():
    if request.method == "POST":

        title = request.form['title']
        intro = request.form['intro']
        text = request.form['text']

        article = Article(title=title, intro=intro, text=text)
        # opt = T_division.query.order_by(T_division.division_name.desc()).all()
        try:
            db.session.add(article)
            db.session.commit()
            return redirect('/posts')
        except:
            return "При добавление статьи произошла ошибка"
    else:
        return render_template("create-article.html")


@app.route('/create-lic', methods=['POST', 'GET'])
@login_required
def create_lic():
    option = T_division.query.order_by(T_division.division_name.desc()).all()

    if request.method == "POST":

        fio = request.form['fio']
        computer_name = request.form['computer_name']
        license_code = request.form['license_code']
        contract = request.form['contract']
        comment = request.form['comment']
        code = request.form['code']
        div_id = request.form['list_name']
        # intro = request.form['intro']
        # text = request.form['text']

        # article = Article(title=title, intro=intro, text=text)
        # opt = T_division.query.order_by(T_division.division_name.desc()).all()
        lic = T_license(fio=fio, computer_name=computer_name, license_code=license_code, contract=contract,
                        comment=comment, code=code, div_id=int(div_id))
        try:
            db.session.add(lic)
            db.session.commit()
            return redirect('/license')
        except:
            return "При добавление записи произошла ошибка"
    else:
        return render_template("create-lic.html", option=option)

    # articles = Article.query.order_by(Article.date.desc()).all()
    # return render_template("posts.html", articles=articles)


@app.route('/create-lic2', methods=['POST', 'GET'])
@login_required
def create_lic2():
    option = T_division.query.order_by(T_division.division_name.asc()).all()
    post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
    sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()

    if request.method == "POST":

        fio = request.form['fio'] # ФИО
        div_id = request.form['division'] # Подразделение
        post_id = request.form['post'] # Должность
        license_code = request.form['lic_code'] # Лицензия
        license_owner_id = request.form['lic_owner_id'] # Где оригинал
        computer_name = request.form['name_comp'] # Имя компьютера
        comment = request.form['comment'] # Примечание
        contract = request.form['contract'] # Контракт
        code = request.form['code'] # Код

        par = request.form['submitType']

        print(request.form.getlist('list'))

        # intro = request.form['intro']
        # text = request.form['text']

        # article = Article(title=title, intro=intro, text=text)
        # opt = T_division.query.order_by(T_division.division_name.desc()).all()
        lic = T_license(fio=fio, div_id=int(div_id), post_id=int(post_id), license_code=license_code, license_owner_id=int(license_owner_id), computer_name=computer_name,
                        comment=comment, code=code, contract=contract)
        sp = request.form.getlist('list')

        us = current_user.id

        try:
            db.session.add(lic)
            db.session.flush()

            db.session.refresh(lic)
            # print(lic.id)
            log = T_log(LICENSE_ID=lic.id, EVENT_TYPE=1, COMMENT=comment, USER_ID=us)

            for s in request.form.getlist('list'):
                print(s)
                sys = TC_LICENSE_SYSTEM(SYSTEM_ID=s, LICENSE_ID=lic.id)
                db.session.add(sys)

            db.session.add(log)
            db.session.commit()

            if par == 'Buy':
                return redirect('/license/' + str(lic.id) + '/update')
            else:
                return redirect('/license')

        except:
            return "При добавление статьи произошла ошибка"
    else:
        return render_template("create-lic2.html", option=option, post_sp=post_sp, sys_sp=sys_sp)


# Создание акта передачи
@app.route('/create-act', methods=['POST', 'GET'])
@login_required
def create_act():
    option = T_division.query.order_by(T_division.division_name.asc()).all()
    post_sp = Post.query.order_by(Post.POST_NAME.asc()).all()
    name_skzi = Name_SKZI.query.order_by(Name_SKZI.n_skzi).all()
    # sys_sp = TS_SYSTEM.query.order_by(TS_SYSTEM.SYSTEM_NAME.desc()).all()
    employee = Profiles.query.order_by(Profiles.name.asc()).all()
    # Условие отбора новых лицензий для передачи.
    # --- Пока не понятно !!!
    lic_empty = T_license.query.filter(T_license.div_id == 1).all()
    uuid_g = uuid.uuid4()

    if request.method == "POST":

        user_sour = request.form['fio_source'] # сотрудник 1
        user_rec = request.form['fio_receiver'] # сотрудник 2
        name_skzi = request.form['skzi'] #Наименование средств ЭП
        comment = request.form['comment'] # примечание
        # DIV_SOUR_ID = request.form['division'] # организация , где хранится оригинал
        DIVISION_ID = request.form['division'] # Наименование организации получателя (div_id в лицензии)
        date_received = request.form['data']  # Получено дата

        # fio = request.form['fio'] # ФИО
        # div_id = request.form['division'] # Подразделение
        # post_id = request.form['post'] # Должность
        # license_code = request.form['lic_code'] # Лицензия
        # license_owner_id = request.form['lic_owner_id'] # Где оригинал
        # computer_name = request.form['name_comp'] # Имя компьютера
        # comment = request.form['comment'] # Примечание
        # contract = request.form['contract'] # Контракт
        # code = request.form['code'] # Код

        par = request.form['submitType']

        sp = request.form.getlist('key_sp')

        us = current_user.id

        # val = request.form.getlist('key_sp')
        #
        # n = len(val)
        # m = len(request.form.getlist('key_sp'))

        # if len(request.form.getlist('key_sp')):
        if len(request.form.getlist('key_sp')) == 0:
            return "Не выбраны лицензии!"

        try:
            # db.session.add(lic)
            # db.session.flush()
            #
            # db.session.refresh(lic)
            # print(lic.id)
            # log = T_log(LICENSE_ID=lic.id, EVENT_TYPE=1, COMMENT=comment, USER_ID=us)

            for s in request.form.getlist('key_sp'):

                pr_j = Journal_skzi.query.filter(Journal_skzi.LICENSE_ID==s).first()

                if pr_j:
                    return "Лицензия уже передана"

                lic2 = T_license.query.get(s)
                DIV_SOUR_ID = lic2.license_owner_id
                spisok_j = Journal_skzi(user_sour=user_sour, user_rec=user_rec, name_skzi=name_skzi, comment=comment,
                                        DIV_SOUR_ID=DIV_SOUR_ID, DIVISION_ID=DIVISION_ID, date_received=date_received, LICENSE_ID=s, uuid=uuid_g)
                lic2.div_id = DIVISION_ID
                # sys = TC_LICENSE_SYSTEM(SYSTEM_ID=s, LICENSE_ID=lic.id)
                db.session.add(spisok_j)

            # db.session.add(log)
            db.session.commit()
            return redirect('/journal')
            # if par == 'Buy':
            #     return redirect('/license/' + str(lic.id) + '/update')
            # else:
            #     return redirect('/license')

        except:
            return "При добавление статьи произошла ошибка"
    else:
        return render_template("create-act.html", option=option, post_sp=post_sp, employee=employee, name_skzi=name_skzi, lic_empty=lic_empty)



@app.route('/create-key/<int:id>', methods=['POST', 'GET'])
@login_required
def create_key(id):
    spis_div = T_division.query.order_by(T_division.division_name.asc()).all()
    spis_centr = TS_CENTER.query.order_by(TS_CENTER.id.asc()).all()


    if request.method == "POST":
        ID_CENTER = request.form['ud_centr']
        DATE_STOP = request.form['data_stop'] #
        KEY_OWNER_ID = request.form['division'] # Подразделение
        GOST = request.form['gost']
        COMMENT = request.form['rezen']
        ID_LICENSE = id

        us = current_user.id

        sp_key = T_Key(DATE_STOP=DATE_STOP, GOST=GOST, COMMENT=COMMENT, ID_CENTER=ID_CENTER, KEY_OWNER_ID=KEY_OWNER_ID, ID_LICENSE=ID_LICENSE)
        log = T_log(LICENSE_ID=id, EVENT_TYPE=2, COMMENT='', USER_ID=us)

        try:
            db.session.add(sp_key)
            db.session.add(log)

            db.session.commit()
            return redirect('/key/' + str(id))
        except:
            return "При добавление ключа произошла ошибка"
    else:
        return render_template("create-key.html", spis_div=spis_div, spis_centr=spis_centr)


# @app.route("/get-csv/<path:filename>")
# def get_csv(filename):
#     # безопасно соединяем базовый каталог и имя файла
#     safe_path = safe_join(app.config["UPLOAD_FOLDER"], filename)
#     try:
#         return send_file(safe_path, as_attachment=True)
#         # return send_file('static/uploads/helloworld.docx', as_attachment=True)
#     except FileNotFoundError:
#         abort(404)

# Загрузка с сайта файла docx - Акт о передачи
@app.route('/download/<uuid>')
def download(uuid):

    filename= files(uuid)
    # filename = 'test_123.docx'
    # безопасно соединяем базовый каталог и имя файла
    safe_path = safe_join(app.config["UPLOAD_FOLDER"], filename)
    try:
        return send_file(safe_path, as_attachment=True)
    except FileNotFoundError:
        abort(404)


# Загрузка с сайта файла xlsx - журнал средств СКЗИ
@app.route('/download-skzi')
def download_skzi():
    filename= files_exc()
    # filename = 'test_123.docx'
    # безопасно соединяем базовый каталог и имя файла
    safe_path = safe_join(app.config["UPLOAD_FOLDER"], filename)
    try:
        return send_file(safe_path, as_attachment=True)
    except FileNotFoundError:
        abort(404)


@app.route("/login", methods=['POST', "GET"])
def login():
    # if request.method == 'POST':
    #     if len(request.form['email']) > 2:
    #         flash('Сообщение отправлено')
    #     else:
    #         flash('ошибка отпраки')
    #
    # return render_template("Login.html")
    email = request.form.get('email')
    psw = request.form.get('psw')

    if email and psw:
        user = Users.query.filter_by(email=email).first()

        if user and check_password_hash(user.psw, psw):
            login_user(user)
            us = current_user.id
            # next_page = request.args.get('next')
            #
            # return redirect(next_page)

            next_page = request.args.get('next')

            log = T_log(EVENT_TYPE=8,  USER_ID=us)
            try:
                db.session.add(log)
                db.session.commit()
            except:
                return "Произошла ошибка"
                    # return redirect(next_page)
            return redirect(url_for('list'))
        else:
            flash('Логин или пароль не правильны')
    else:
        flash('Пожалуйста введите логин и пароль')

    return render_template('login.html')


# Получение параметра из страницы
# @app.route('/user/<string:name>/<int:id>')
# def user(name, id): # эти параметры передаются в функцию
#     return "User page: " + name + " - " + str(id)


@app.route("/register", methods=("POST", "GET"))
def register():
    if request.method == "POST":
        # здесь должна быть проверка корректности введенных данных
        if len(request.form['name']) > 4 and len(request.form['email']) > 4 and len(request.form['psw']) > 4 and \
                request.form['psw'] == request.form['psw2']:
            try:
                hash = generate_password_hash(request.form['psw'])
                u = Users(email=request.form['email'], psw=hash)
                # Проверка e-mail на дубли
                res = Users.query.filter(Users.email.like(u.email)).count()
                #
                if res == 0:
                    db.session.add(u)
                    db.session.flush()

                    p = Profiles(name=request.form['name'], old=request.form['old'],
                                 city=request.form['city'], user_id=u.id)
                    db.session.add(p)
                    db.session.commit()
                    flash("вы успешно зарегистрированы в БД", "success")
                    return redirect(url_for('login'))
                else:
                    flash("Пользователь с таким e-mail существует в базе", "error")
            except:
                db.session.rollback()
                print("Ошибка добавления в БД")
                flash("Ошибка добавления в БД", "error")
        else:
            flash("Неверно заполнены поля", "error")

    return render_template("register.html", title="Регистрация")


@app.route('/logout', methods=['GET', 'POST'])
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))


@app.after_request
def redirect_to_signin(response):
    if response.status_code == 401:
        return redirect(url_for('login')+'?next='+request.url)

    return response



#     Добавить несколько страниц в документ


#
def files(uuid):

    a = aliased(T_division)
    res = db.session.query(Journal_skzi, T_division, a, T_license, Name_SKZI).join(T_division,
                                                                        Journal_skzi.DIVISION_ID == T_division.id).join(
        a, Journal_skzi.DIV_SOUR_ID == a.id).join(T_license, Journal_skzi.LICENSE_ID == T_license.id).join(Name_SKZI, Journal_skzi.name_skzi==Name_SKZI.id).filter(Journal_skzi.uuid==uuid).all()
    # ---

    listPost = []
    # query = 'select sel_pfio(\'a57d3af9-cff9-488c-9c20-68ace936de70\')'
    val_uuid = '\'' +  uuid +'\''
    # query = 'select sel_pfio(' + val_uuid + ')'
    query = 'select sel_pfio('+val_uuid+')'
    results = db.session.execute(query)

    k=0
    for r in results:
        listPost.append(r[0])
        k=k+1

    # print(listPost[0],'---' ,listPost[1])
    # for r in listPost:
    #     print(r)

    # print(len(res))
    # print(res[0][0].uuid)
    # for r in res:
    #     print(r[3].code)

    # div_recipient = res[0][1].division_name#Организация - получатель
    # name_skzi = res[0][4].n_skzi# наименование средста эп
    # Col = # Количество
    # Список лицензий:
    #
    #
    #
    # Post1 # Должность сотрудника передающего средство СКЗИ
    # FIO1  # ФИО сотрудника передающего средство СКЗИ
    # Post2 # Должность сотрудника принимающего средство
    # FIO2  # сотрудника принимающего средство СКЗИ

    len_res = len(res)
    # ---
    zet = len_res % 2
    # print(zet)
    if zet==0:
        col1 = int(len_res / 2)
        col2 = col1
    elif zet==1:
        col1 = int(len_res // 2 + 1)
        col2 = int(len_res // 2)


    wordDoc = Document("static/uploads/lic_kriptopro.docx")

    table = wordDoc.tables[0]

    cell = table.cell(0, 1)

    table.cell(0, 1).text = res[0][1].division_name # Организация - получатель
    rc = table.cell(0, 1).paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(12)
    table.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # p = head_cells[i].paragraphs[0]
    # название колонки
    # p.add_run(item).bold = True
    # # выравниваем посередине
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.cell(1, 1).text = res[0][4].n_skzi # Наименование средста эп
    rc = table.cell(1, 1).paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(12)
    table.cell(1, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table.cell(2, 1).text = str(len_res) # Количество
    rc = table.cell(2, 1).paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(12)
    table.cell(2, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = wordDoc.tables[2]
    table.cell(0, 0).text = listPost[0]# Должность 1
    rc = table.cell(0, 0).paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(12)
    # table.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table.cell(0, 2).text = listPost[2]# ФИО 1
    rc = table.cell(0, 2).paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(12)
    table.cell(0, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table.cell(2, 0).text = listPost[1]#  Должность 2
    rc = table.cell(2, 0).paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(12)

    table.cell(2, 2).text = listPost[3]# ФИО 1
    rc = table.cell(2, 2).paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(12)
    # table.cell(2, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # for table in wordDoc.tables:
    # for row in table.rows:
    #     for cell in row.cells:
    #         print(cell.text)
    # print(len(wordDoc.tables))
    row = wordDoc.tables[1].rows[0]

    # for row in wordDoc.tables[1].table.rows:
    #     for cell in row.cells:
    #         print(cell.text)

    for i in range(col1):
        # print('1--- ', i)
        if i==0:
            row.cells[0].text = res[i][3].code
            # rc = row.paragraphs[0].runs[0]
            row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # table.cell(0, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            wordDoc.tables[1].add_row()
            row = wordDoc.tables[1].rows[i]
            row.cells[0].text = res[i][3].code
            row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    for i in range(col2):
        # print('2--- ', i)
        row = wordDoc.tables[1].rows[i]
        row.cells[1].text = res[col1 + i][3].code
        row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # wordDoc.save("static/uploads/test_123.docx")
    # m = UPLOAD_FOLDER + uuid + '.docx'
    n ='static/uploads/' + uuid
    name_file = n + '.docx'
    wordDoc.save(name_file)
    return (uuid + '.docx')
#

# Вспомогательная ф-я для files_exc
def set_border(ws, cell_range):
    thin = Side(border_style="thin", color="000000")
    for row in ws[cell_range]:
        for cell in row:
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
            cell.alignment = Alignment(horizontal='center', wrap_text=True)


# функция для создания файла выгрузки
def files_exc():
    users_o = aliased(Users)
    users_i = aliased(Users)
    users_d = aliased(Users)
    prof_o = aliased(Profiles)
    prof_i = aliased(Profiles)
    prof_d = aliased(Profiles)

    sp = db.session.query(T_skzi, Name_SKZI, T_division, users_i, users_o, users_d, prof_o, prof_i, prof_d).filter(
        T_skzi.name_skzi == Name_SKZI.id).filter(T_skzi.div_sours == T_division.id).filter(
        T_skzi.fio_owner == users_o.id).filter(users_o.id == prof_o.user_id).filter(
        T_skzi.fio_instal == users_i.id).filter(users_i.id == prof_i.user_id).filter(
        T_skzi.fio_destr == users_d.id).filter(users_d.id == prof_d.user_id).order_by(T_skzi.id.asc()).all()

    wb = load_workbook('static/uploads/temp_jornal.xlsx')
    # print(wb.sheetnames)
    ws = wb["ЖурналСКЗИ"]

    # ws = wb["Mysheet2"]

    # data = [ [row*col for col in range(1, 10)] for row in range(1, 31)]
    #
    data = []
    # for row in range(1, 31):
    #     for col in range(1, 10):
    #         data.append(row*col)

    cx = 0

    data = [['0', el[0].n_account, el[1].n_skzi, el[0].n_serial, el[0].n_instances, el[2].division_name, el[0].sopr_pis,
             el[6].name, el[0].date_rec.strftime('%d-%m-%Y'),
             el[7].name, el[0].date_con.strftime('%d-%m-%Y'), el[0].n_hardware, el[0].data_destr.strftime('%d-%m-%Y'),
             el[8].name, el[0].n_destr, el[0].note] for el in sp]
    #
    for i in range(0, len(data)):
        data[i][0] = i+1

    # l = len (sp)
    # for row in data:
    #     ws.append(row)
    r = 0
    c = 0

    # print('---', data[0][15])

    for row in range(4, len(sp)+4):
        for col in range(1, 17):
            # print(value)
            cell = ws.cell(row=row, column=col)
            # print(row, col, '--', r, '--', c)
            value = data[r][c]
            cell.value = value
            c = c+1
        r = r+1
        c = 0

    # row = ws.max_row
    #
    # thin = Side(border_style="thin", color="000000")
    #
    # for col in range(1, ws.max_column):
    #     cell = ws.cell(row=row, column=col)
    #
    # cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

    rang = 'A4'+':'+'P'+str(len(sp)+3)
    set_border(ws, rang)

    # print(c)
    # r = 1
    # c = 1
    # for row in range(3, 7):
    # for row in sp:
    #     for col in range(1, 17):
    #         value = str(row) + str(col)
    #         cell = ws.cell(row = row, column = col)
    #         cell.value = value
    #         c = c +1
    #     r = r +1
    #     c = 0
    name = 'static/uploads/journal.xlsx'
    wb.save(name)

    return 'journal.xlsx'



if __name__ == "__main__":
    app.run(debug=True)
